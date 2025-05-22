import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# ---------- FUNCIONES ----------
def generar_titular_diario(row):
    return f"Titular automático para el {row['fecha'].strftime('%d-%m-%Y')}"

def generar_noticia_breve_para_fecha(fecha, df):
    return f"Resumen breve del tiempo para el {fecha.strftime('%d-%m-%Y')}."

def generar_noticia_con_comparativas(fecha_str, df):
    return f"Comparativa del tiempo con otros días en base al {fecha_str}."

# ---------- INTERFAZ STREAMLIT ----------
st.title("📰 Generador de Noticias del Tiempo")

# Subida del archivo Excel
archivo = st.file_uploader("📁 Sube tu archivo Excel", type=["xls", "xlsx"])
if archivo:
    df = pd.read_excel(archivo)
    
    # Convertir la columna 'fecha' a tipo datetime
    if not pd.api.types.is_datetime64_any_dtype(df['fecha']):
        df['fecha'] = pd.to_datetime(df['fecha'])

    # Selección de fecha
    fecha_input = st.date_input("📆 Selecciona una fecha para generar la noticia")

    # Botón para generar
    if st.button("🛠️ Generar Noticia"):
        try:
            row = df[df['fecha'] == pd.to_datetime(fecha_input)].iloc[0]
            
            # Generar textos
            titular = generar_titular_diario(row)
            noticia_breve = generar_noticia_breve_para_fecha(fecha_input, df)
            noticia_comparativa = generar_noticia_con_comparativas(fecha_input.strftime('%d-%m-%Y'), df)

            # Crear documento
            doc = Document()
            doc.add_heading(titular, level=1)
            doc.add_paragraph("")
            p_breve = doc.add_paragraph(style='List Bullet')
            run_breve = p_breve.add_run(noticia_breve)
            run_breve.bold = True
            doc.add_paragraph("")
            doc.add_paragraph(noticia_comparativa)

            # Guardar en memoria para descarga
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            nombre_archivo = f"noticia_tiempo_{fecha_input.strftime('%Y_%m_%d')}.docx"
            st.success("✅ Documento generado correctamente.")
            st.download_button("📥 Descargar documento Word", buffer, file_name=nombre_archivo,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except IndexError:
            st.error("❌ No se encontró información para la fecha seleccionada.")
